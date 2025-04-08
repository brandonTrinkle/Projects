import docx

def main():
    # Create a new Document
    doc = docx.Document()
    
    # Title
    doc.add_heading("MAT 243 – Discrete Math Structures: Exam 2 Comprehensive Cheat Sheet", 0)
    
    # I. Proof Techniques
    doc.add_heading("I. Proof Techniques", level=1)
    
    doc.add_heading("1. Direct Proof", level=2)
    doc.add_paragraph("• Assume the hypothesis holds for an arbitrary element (state the domain).")
    doc.add_paragraph("• Use definitions, algebraic manipulation, and known theorems to deduce the conclusion.")
    doc.add_paragraph("• Example: To show “if x is even, then x + 3 is odd”, let x = 2k and write:")
    doc.add_paragraph("      x + 3 = 2k + 3 = 2(k + 1) + 1 (odd number).")
    
    doc.add_heading("2. Proof by Contraposition", level=2)
    doc.add_paragraph("• To prove 'if P then Q', prove instead 'if not Q then not P.'")
    doc.add_paragraph("• Assume ¬Q and show that ¬P follows.")
    doc.add_paragraph("• Useful when the direct proof is not straightforward.")
    
    doc.add_heading("3. Proof by Contradiction", level=2)
    doc.add_paragraph("• Assume the negation of the statement to be proved.")
    doc.add_paragraph("• Deduce a logical contradiction (e.g., violation of a known property).")
    doc.add_paragraph("• Conclude the original statement must be true.")
    
    doc.add_heading("4. Proofs with Quantifiers", level=2)
    doc.add_paragraph("• Universal Statements (∀x P(x)):")
    doc.add_paragraph("   – Prove by letting x be an arbitrary element and showing P(x) holds.")
    doc.add_paragraph("   – Disprove by finding a counterexample: ∃x such that ¬P(x).")
    doc.add_paragraph("• Existential Statements (∃x P(x)):")
    doc.add_paragraph("   – Prove by providing a specific example x₀ with P(x₀) true.")
    doc.add_paragraph("   – Disprove by showing ∀x ¬P(x).")
    doc.add_paragraph("• Order matters: ∀x∃y P(x,y) vs. ∃y∀x P(x,y).")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # II. Sets and Set Operations
    doc.add_heading("II. Sets and Set Operations", level=1)
    
    doc.add_heading("1. Basic Definitions", level=2)
    doc.add_paragraph("• A set is an unordered collection of distinct objects (e.g., A = {1, 2, 3}).")
    doc.add_paragraph("• The empty set, ∅, contains no elements, while {∅} is a singleton (one element).")
    
    doc.add_heading("2. Power Set", level=2)
    doc.add_paragraph("• P(S) is the set of all subsets of S.")
    doc.add_paragraph("   – Example: If S = {1, 2}, then P(S) = {∅, {1}, {2}, {1,2}}.")
    
    doc.add_heading("3. Subsets", level=2)
    doc.add_paragraph("• A ⊆ B means every element of A is in B.")
    doc.add_paragraph("• A ⊂ B (proper subset) means A ⊆ B but A ≠ B (no set is a proper subset of itself).")
    
    doc.add_heading("4. Set Operations", level=2)
    doc.add_paragraph("• Union (A ∪ B): Elements in A or B (or both). Example: {1,2,3} ∪ {3,4,5} = {1,2,3,4,5}.")
    doc.add_paragraph("• Intersection (A ∩ B): Common elements. Example: {1,2,3} ∩ {3,4,5} = {3}.")
    doc.add_paragraph("• Difference (A − B): Elements in A that are not in B. Example: {1,2,3} − {2,3} = {1}.")
    doc.add_paragraph("• Complement: Relative to a universal set U. Example: If S = (2,5) ⊆ ℝ, then Sᶜ = (−∞,2] ∪ [5,∞).")
    doc.add_paragraph("• Symmetric Difference (A △ B): Elements in exactly one of A or B. Example: {1,2,3} △ {3,4,5} = {1,2,4,5}.")
    doc.add_paragraph("• Cartesian Product (A × B): Set of all ordered pairs (a, b) with a ∈ A and b ∈ B.")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # III. Functions
    doc.add_heading("III. Functions", level=1)
    
    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph("• A function f: A → B assigns every element x ∈ A exactly one element f(x) ∈ B.")
    doc.add_paragraph("• Domain: The set A of inputs.")
    doc.add_paragraph("• Codomain: The set B of possible outputs.")
    doc.add_paragraph("• Range: {f(x) | x ∈ A} (the actual outputs).")
    
    doc.add_heading("2. Properties", level=2)
    doc.add_paragraph("• Injective (one-to-one): f(x₁) = f(x₂) implies x₁ = x₂.")
    doc.add_paragraph("• Surjective (onto): For every y ∈ B, there exists x ∈ A such that f(x) = y.")
    doc.add_paragraph("• Bijective: Both injective and surjective; f has an inverse.")
    doc.add_paragraph("• Well-Defined: Every input maps to an element within the codomain.")
    
    doc.add_heading("3. Images and Preimages", level=2)
    doc.add_paragraph("• For X ⊆ A, f(X) = {f(x) | x ∈ X}.")
    doc.add_paragraph("• For Y ⊆ B, f⁻¹(Y) = {x ∈ A such that f(x) ∈ Y}.")
    
    doc.add_heading("4. Review Notes", level=2)
    doc.add_paragraph("• A function is surjective if its range equals its codomain.")
    doc.add_paragraph("• In finite sets of equal size, injectivity implies surjectivity.")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # IV. Floor and Ceiling Functions
    doc.add_heading("IV. Floor and Ceiling Functions", level=1)
    
    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph("• Floor function ⌊x⌋: The greatest integer ≤ x.")
    doc.add_paragraph("• Ceiling function ⌈x⌉: The smallest integer ≥ x.")
    doc.add_paragraph("• Inequalities: For any real x, x − 1 < ⌊x⌋ ≤ x and x ≤ ⌈x⌉ < x + 1.")
    
    doc.add_heading("2. Example Problem", level=2)
    doc.add_paragraph("• Solve: 2 ≤ ⌈2x + 5⌉ < 5. (Rewrite in terms of 2x + 5 and solve for x.)")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # V. Sequences, Summation, and Sigma Notation
    doc.add_heading("V. Sequences, Summation, and Sigma Notation", level=1)
    
    doc.add_heading("1. Sequences", level=2)
    doc.add_paragraph("• A sequence is an ordered list denoted {aₙ}, where n is an index.")
    doc.add_paragraph("• Indexing: Zero-based (a₀ is the first term) or one-based (a₁ is the first term).")
    
    doc.add_heading("2. Types of Sequences", level=2)
    doc.add_paragraph("• Arithmetic: aₙ = a + n·d, where d is the constant difference.")
    doc.add_paragraph("• Geometric: aₙ = a · qⁿ, where q is the constant ratio.")
    
    doc.add_heading("3. Recursive Definitions", level=2)
    doc.add_paragraph("• A sequence can be defined recursively (e.g., aₙ = aₙ₋₁ + 2 with a₀ = 0).")
    
    doc.add_heading("4. Summation and Sigma Notation", level=2)
    doc.add_paragraph("• Express sums using Σ notation.")
    doc.add_paragraph("• Practice index shifts and writing closed forms.")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # VI. Additional Review Exam Concepts
    doc.add_heading("VI. Additional Review Exam Concepts", level=1)
    
    doc.add_heading("1. Function Properties (Review Q1–Q2)", level=2)
    doc.add_paragraph("• Understand domain, codomain, and range.")
    doc.add_paragraph("• For f: A → B with |A| = |B|, injectivity implies surjectivity.")
    doc.add_paragraph("• If f is not surjective, then its range is a proper subset of the codomain.")
    
    doc.add_heading("2. Proof by Contradiction (Review Q3)", level=2)
    doc.add_paragraph("• For “if x is rational and y is irrational then x+y is irrational”, assume:")
    doc.add_paragraph("    - x is rational, y is irrational, and x+y is rational, then derive a contradiction.")
    
    doc.add_heading("3. Set Operations (Review Q6–Q7)", level=2)
    doc.add_paragraph("• Find power sets and Cartesian products (e.g., power set of A = {1, {a, t}}, compute A × A).")
    doc.add_paragraph("• Understand membership vs. subset:")
    doc.add_paragraph("    - {1,2} ⊆ {1,2,3,A,c} is true.")
    doc.add_paragraph("    - {1} ∈ {1,2,3,A,c} is true.")
    doc.add_paragraph("    - ∅ ⊆ S is always true; ∅ ∈ S only if ∅ is an element of S.")
    
    doc.add_heading("4. Quantified Statements (Review Q8)", level=2)
    doc.add_paragraph("• Practice proving:")
    doc.add_paragraph("    a. ∃x ∀y (5y − xy = y) for all real numbers.")
    doc.add_paragraph("    b. ∀x ∃y (y is odd and x − 1 < y ≤ x + 1) for positive integers.")
    doc.add_paragraph("    c. ∀x ∃y ∃z (y ≠ x + z) for all integers.")
    
    doc.add_heading("5. Floor/Ceiling & Function Images (Review Q9–Q10)", level=2)
    doc.add_paragraph("• Evaluate f([−1,5]) when f is the ceiling function.")
    doc.add_paragraph("• Disprove: For all real x, y, ⌊x + y⌋ = ⌊x⌋ + ⌊y⌋.")
    
    doc.add_heading("6. Function Analysis (Review Q11)", level=2)
    doc.add_paragraph("• For f: (–2,2) → (–4,0] defined by f(x) = –x², check if f is increasing/decreasing, injective, surjective, etc.")
    
    doc.add_heading("7. Summation Problems (Review Q12–Q15)", level=2)
    doc.add_paragraph("• Practice index shifts in summation.")
    doc.add_paragraph("• Write sums in sigma notation and find closed forms.")
    doc.add_paragraph("• Example: Express and evaluate the sum 26 + 30 + 34 + … + 554.")
    
    doc.add_paragraph("---------------------------------------------------------------------")
    
    # VII. Final Reminders
    doc.add_heading("VII. Final Reminders", level=1)
    doc.add_paragraph("• Always state your assumptions and domain clearly in proofs.")
    doc.add_paragraph("• Justify every logical step.")
    doc.add_paragraph("• Use proper notation: distinguish between ∈ and ⊆, and ordered pairs (a, b) vs. sets {a, b}.")
    doc.add_paragraph("• Practice writing proofs with quantifiers and review common pitfalls.")
    doc.add_paragraph("• Use visual aids (Venn diagrams, number lines) to understand sets and intervals.")
    doc.add_paragraph("• Revisit homework (Week 3) and the Review Exam 2 document to cover all topics.")
    
    doc.add_paragraph("=====================================================================")
    doc.add_paragraph("Good luck on Exam 2 – Study hard and review these concepts thoroughly!")
    doc.add_paragraph("=====================================================================")
    
    # Save the document as a .docx file
    doc.save("Exam2StudyGuide.docx")
    print("Cheat sheet saved to Exam2StudyGuide.docx")

if __name__ == "__main__":
    main()
