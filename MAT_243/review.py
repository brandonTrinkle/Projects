from docx import Document

# Create a new Document
document = Document()

# Title
document.add_heading("Study Guide for MAT 243 Review Exam 1", 0)

# Section: Symbols and Logical Operators
document.add_heading("Symbols and Logical Operators", level=1)
document.add_paragraph("• p, q, r, a: Propositional variables representing statements.")
document.add_paragraph("• → : Implies (if ... then).")
document.add_paragraph("• ∧ : Logical AND (conjunction).")
document.add_paragraph("• ∨ : Logical OR (disjunction).")
document.add_paragraph("• ¬ : Logical NOT (negation).")
document.add_paragraph("• ∀ : Universal quantifier (for all).")
document.add_paragraph("• ∃ : Existential quantifier (there exists).")
document.add_paragraph("• ↔ : Biconditional (if and only if).")

# Section 1: Proposition
document.add_heading("1. Proposition: Is 'Frederic Gauss is the greatest mathematician of all time' a Proposition?", level=1)
document.add_paragraph(
    "Key Idea: In many logic texts, a proposition is defined as a statement that is either true or false. "
    "A proper proposition must have an objectively verifiable truth value."
)
document.add_paragraph(
    "According to your professor, the statement is an opinion that cannot be objectively verified, and therefore, "
    "it is not considered a proposition."
)

# Section 2: Important Logical Equivalences
document.add_heading("2. Important Logical Equivalences", level=1)

# a. Conditional as a Disjunction
document.add_heading("a. Conditional as a Disjunction", level=2)
document.add_paragraph("Expression: p → q ≡ ¬p ∨ q")
document.add_paragraph("Plain Language: 'If p then q' means either p is false or q is true.")

# b. Contrapositive
document.add_heading("b. Contrapositive of a Conditional Statement", level=2)
document.add_paragraph("Expression: p → q ≡ ¬q → ¬p")
document.add_paragraph("Plain Language: If 'p implies q' is true, then 'if q is false, then p must be false' is also true.")

# c. Negation of a Conditional
document.add_heading("c. Negation of a Conditional Statement", level=2)
document.add_paragraph("Expression: ¬(p → q) ≡ p ∧ ¬q")
document.add_paragraph("Plain Language: Saying 'it is not true that if p then q' is equivalent to 'p is true and q is false'.")

# d. De Morgan's Identities
document.add_heading("d. De Morgan’s Identities", level=2)
document.add_paragraph("For Conjunction: ¬(p ∧ q) ≡ ¬p ∨ ¬q")
document.add_paragraph("For Disjunction: ¬(p ∨ q) ≡ ¬p ∧ ¬q")
document.add_paragraph("Plain Language: These identities explain how negation distributes over 'and' and 'or'.")

# Section 3: Proving an Equivalence
document.add_heading("3. Prove that (p ∧ q) → r ≡ (p → r) ∨ (q → r)", level=1)
document.add_paragraph("Goal: Prove that (p ∧ q) → r is logically equivalent to (p → r) ∨ (q → r).")
document.add_paragraph("Step-by-Step:")
document.add_paragraph("1. Start with (p ∧ q) → r.")
document.add_paragraph("2. Rewrite as a disjunction: (p ∧ q) → r ≡ ¬(p ∧ q) ∨ r.")
document.add_paragraph("3. Apply De Morgan’s Law: ¬(p ∧ q) ≡ ¬p ∨ ¬q, so the expression becomes ¬p ∨ ¬q ∨ r.")
document.add_paragraph("4. Note that p → r ≡ ¬p ∨ r and q → r ≡ ¬q ∨ r.")
document.add_paragraph("5. Therefore, (p → r) ∨ (q → r) ≡ (¬p ∨ r) ∨ (¬q ∨ r), which simplifies (by associativity/commutativity) to ¬p ∨ ¬q ∨ r.")
document.add_paragraph("6. Conclusion: Since both sides simplify to the same expression, the equivalence holds.")

# Section 4: Negating Statements
document.add_heading("4. Express the Negation of the Following Statements", level=1)

# a.
document.add_heading("a. Negate: ∀x ∃y (xy ≤ 0 → (-1 ≤ y < 0))", level=2)
document.add_paragraph("Steps:")
document.add_paragraph("• Negate the entire statement: ¬[∀x ∃y (xy ≤ 0 → (-1 ≤ y < 0))]")
document.add_paragraph("• Change the quantifiers: ≡ ∃x ∀y ¬(xy ≤ 0 → (-1 ≤ y < 0))")
document.add_paragraph("• Negate the conditional using ¬(A → B) ≡ A ∧ ¬B: ¬(xy ≤ 0 → (-1 ≤ y < 0)) ≡ (xy ≤ 0) ∧ ¬(-1 ≤ y < 0)")
document.add_paragraph("• Final Form: ∃x ∀y [(xy ≤ 0) ∧ ¬(-1 ≤ y < 0)]")

# b.
document.add_heading("b. Negate: 'Every planet in this solar system revolves around the Sun.'", level=2)
document.add_paragraph("Steps:")
document.add_paragraph("• Express in logical form: ∀x (P(x) → R(x)), where")
document.add_paragraph("  P(x): 'x is a planet in this solar system'")
document.add_paragraph("  R(x): 'x revolves around the Sun'")
document.add_paragraph("• Negate and change quantifiers: ¬∀x (P(x) → R(x)) ≡ ∃x ¬(P(x) → R(x))")
document.add_paragraph("• Negate the conditional: ¬(P(x) → R(x)) ≡ P(x) ∧ ¬R(x)")
document.add_paragraph("• Final Form: ∃x [P(x) ∧ ¬R(x)]")

# c.
document.add_heading("c. Negate: 'Some people in Arizona don’t like to drive on highways.'", level=2)
document.add_paragraph("Steps:")
document.add_paragraph("• Logical form: ∃x [A(x) ∧ ¬D(x)], where")
document.add_paragraph("  A(x): 'x is a person in Arizona'")
document.add_paragraph("  D(x): 'x likes to drive on highways'")
document.add_paragraph("• Negate the statement: ¬∃x [A(x) ∧ ¬D(x)] ≡ ∀x ¬[A(x) ∧ ¬D(x)]")
document.add_paragraph("• Apply De Morgan’s Law: ¬[A(x) ∧ ¬D(x)] ≡ ¬A(x) ∨ D(x)")
document.add_paragraph("• Final Form: ∀x [¬A(x) ∨ D(x)]")

# Section 5: Translating with 'Necessary', 'Sufficient', 'Only If', and 'Unless'
document.add_heading("5. Translating with 'Necessary', 'Sufficient', 'Only If', and 'Unless'", level=1)
document.add_paragraph("Original Statement: 'If there are no clouds, then I can see the stars.'")
document.add_paragraph("Let C = 'there are clouds' and S = 'I can see the stars'.")
document.add_paragraph("Direct form: ¬C → S")
document.add_paragraph("Alternate phrasings:")
document.add_paragraph("• Only if: 'I can see the stars only if there are no clouds.' (Implies: S → ¬C)")
document.add_paragraph("• Necessary: 'For me to see the stars, it is necessary that there are no clouds.'")
document.add_paragraph("• Sufficient: 'The absence of clouds is sufficient for me to see the stars.'")
document.add_paragraph("• Unless: 'I can see the stars unless there are clouds.' (Interpreted as: if there are clouds then I cannot see the stars, i.e., C → ¬S)")

# Section 6: Expressing Statements in "If... then" Form
document.add_heading("6. Expressing Statements in 'If... then' Form", level=1)

document.add_heading("a. 'I get my work done on time only if I don’t procrastinate.'", level=2)
document.add_paragraph("Let W = 'work is done on time' and P = 'I procrastinate'.")
document.add_paragraph("Translation: W → ¬P")

document.add_heading("b. 'It is necessary to have wind to sail across the lake.'", level=2)
document.add_paragraph("Let S = 'sail across the lake' and W = 'have wind'.")
document.add_paragraph("Translation: S → W")

document.add_heading("c. 'To be able to vote in the upcoming election it is sufficient to have an American passport.'", level=2)
document.add_paragraph("Let V = 'able to vote' and P = 'have an American passport'.")
document.add_paragraph("Translation: P → V")

document.add_heading("d. 'Unless we live more economically we are going to use up all our resources.'", level=2)
document.add_paragraph("Let E = 'live more economically' and U = 'use up all resources'.")
document.add_paragraph("Translation: ¬E → U")

# Section 7: Expressing Conditionals as Disjunctions
document.add_heading("7. Expressing Conditionals as Disjunctions", level=1)

document.add_heading("a. Rewrite ¬p → ¬q", level=2)
document.add_paragraph("Using the equivalence: p → q ≡ ¬p ∨ q")
document.add_paragraph("Replace p with ¬p and q with ¬q:")
document.add_paragraph("¬p → ¬q ≡ ¬(¬p) ∨ ¬q ≡ p ∨ ¬q")

document.add_heading("b. Rewrite a → ¬p", level=2)
document.add_paragraph("Using the equivalence:")
document.add_paragraph("a → ¬p ≡ ¬a ∨ ¬p")

# Section 8: Expressing Disjunction as a Conditional
document.add_heading("8. Expressing Disjunction as a Conditional", level=1)
document.add_paragraph("Express p ∨ q using a conditional:")
document.add_paragraph("p ∨ q ≡ ¬p → q")
document.add_paragraph("Plain Language: 'If not p then q' is equivalent to 'p or q.'")

# Section 9: Expressing Conjunction with Conditional and Negation
document.add_heading("9. Expressing Conjunction Using Conditional and Negation", level=1)
document.add_paragraph("Express p ∧ ¬q:")
document.add_paragraph("Observation: The negation of a conditional yields a conjunction, i.e., ¬(p → q) ≡ p ∧ ¬q")
document.add_paragraph("Thus, p ∧ ¬q ≡ ¬(p → q)")

# Section 10: Expressing ¬p ∧ ¬q Using Conditional and Negation
document.add_heading("10. Expressing ¬p ∧ ¬q Using Conditional and Negation", level=1)
document.add_paragraph("Steps:")
document.add_paragraph("1. Start with De Morgan’s Law: ¬p ∧ ¬q ≡ ¬(p ∨ q)")
document.add_paragraph("2. Express the disjunction as a conditional: p ∨ q ≡ ¬p → q")
document.add_paragraph("3. Then, ¬(p ∨ q) ≡ ¬(¬p → q)")
document.add_paragraph("Conclusion: ¬p ∧ ¬q ≡ ¬(¬p → q)")

# Section 11: Evaluating Arguments
document.add_heading("11. Evaluating Arguments", level=1)

document.add_heading("Argument 1: Valid", level=2)
document.add_paragraph("Premises:")
document.add_paragraph("1. Every bug is an insect: ∀x (B(x) → I(x))")
document.add_paragraph("2. Elephants are not insects: ∀x (E(x) → ¬I(x))")
document.add_paragraph("3. Trombone is an elephant: E(Trombone)")
document.add_paragraph("Conclusion: ¬B(Trombone)")
document.add_paragraph("Explanation: Since every bug must be an insect and Trombone (an elephant) is not an insect, it cannot be a bug.")

document.add_heading("Argument 2: Invalid", level=2)
document.add_paragraph("Premises:")
document.add_paragraph("1. Every bug is an insect: ∀x (B(x) → I(x))")
document.add_paragraph("2. Wizzy-Fizz is an insect: I(Wizzy-Fizz)")
document.add_paragraph("Conclusion: B(Wizzy-Fizz)")
document.add_paragraph("Explanation: Although every bug is an insect, not every insect is a bug. This mistake is known as affirming the consequent.")

# Final Tips Section
document.add_heading("Final Tips", level=1)
document.add_paragraph("• Definitions matter – always use the definitions your course provides.")
document.add_paragraph("• Practice converting statements using logical equivalences.")
document.add_paragraph("• Write each step in proofs clearly so you can follow the logic.")
document.add_paragraph("• Identify logical fallacies by checking if conclusions follow from the premises.")

# Save the document as StudyGuide.docx
document.save("StudyGuide.docx")
print("StudyGuide.docx has been created successfully!")
